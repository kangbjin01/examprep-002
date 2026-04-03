#!/usr/bin/env python3
"""
SAT 문제 임포트 스크립트
Usage: python3 scripts/import-sat.py <문제docx경로> <해설docx경로> <회차이름>
Example: python3 scripts/import-sat.py "/path/to/문제.docx" "/path/to/해설.docx" "2025-sep-cookingsat"
"""

import sys
import os
import re
import json
import urllib.request
import urllib.error

POCKETBASE_URL = os.environ.get("POCKETBASE_URL", "http://127.0.0.1:8090")
PB_ADMIN_EMAIL = os.environ.get("PB_ADMIN_EMAIL", "admin@test.com")
PB_ADMIN_PASS = os.environ.get("PB_ADMIN_PASS", "Pass1234!")


def get_admin_token():
    data = json.dumps({
        "identity": PB_ADMIN_EMAIL,
        "password": PB_ADMIN_PASS,
    }).encode("utf-8")
    req = urllib.request.Request(
        f"{POCKETBASE_URL}/api/collections/_superusers/auth-with-password",
        data=data,
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    with urllib.request.urlopen(req) as resp:
        return json.loads(resp.read())["token"]


def parse_questions_docx(path):
    import docx
    doc = docx.Document(path)
    full_text = '\n'.join([p.text.strip() for p in doc.paragraphs if p.text.strip()])

    questions = []
    current_section = None

    parts = re.split(r'(Section\s+\d+\s+Module\s+\d+\s*:\s*.+)', full_text)

    for part in parts:
        sec_match = re.match(r'Section\s+(\d+)\s+Module\s+(\d+)\s*:\s*(.+)', part.strip())
        if sec_match:
            current_section = {
                'section': int(sec_match.group(1)),
                'module': int(sec_match.group(2)),
                'type': sec_match.group(3).strip()
            }
            continue

        if current_section is None:
            continue

        q_parts = re.split(r'(?=Question\s+\d+)', part)
        for qp in q_parts:
            qp = qp.strip()
            if not qp:
                continue
            q_num_match = re.match(r'Question\s+(\d+)', qp)
            if not q_num_match:
                continue

            q_num = int(q_num_match.group(1))
            q_body = qp[q_num_match.end():].strip()

            choices = []
            choice_pattern = r'([A-D])\)\s*(.+?)(?=\n[A-D]\)|$)'
            choice_matches = list(re.finditer(choice_pattern, q_body, re.DOTALL))

            if choice_matches:
                q_text = q_body[:choice_matches[0].start()].strip()
                for cm in choice_matches:
                    choices.append(cm.group(2).strip())
            else:
                q_text = q_body

            # Separate passage from question prompt
            passage = None
            question_text = q_text

            prompt_patterns = [
                r'(Which choice (?:best |most )?(?:completes|describes|states|explains).+)',
                r'(Based on the texts?.+)',
                r'(What is the (?:main|most).+)',
                r'(According to the .+)',
            ]
            for pp in prompt_patterns:
                m = re.search(pp, q_text, re.DOTALL)
                if m:
                    passage = q_text[:m.start()].strip() or None
                    question_text = m.group(1).strip()
                    break

            questions.append({
                'section_num': current_section['section'],
                'module_num': current_section['module'],
                'section_type': current_section['type'],
                'question_number': q_num,
                'passage': passage,
                'question': question_text,
                'choices': choices,
            })

    return questions


def parse_explanation_docx(path):
    import docx
    doc = docx.Document(path)
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    results = {}
    current_section = None
    current_module = None
    answer_counter = 0

    i = 0
    while i < len(lines):
        line = lines[i]

        sec_match = re.match(r'Section\s+(\d+)\s+Module\s+(\d+)\s*:\s*(.+)', line)
        if sec_match:
            current_section = int(sec_match.group(1))
            current_module = int(sec_match.group(2))
            answer_counter = 0
            i += 1
            continue

        q_match = (re.match(r'(\d+)번 문제', line) or
                   re.match(r'Question\s+(\d+)', line) or
                   re.match(r'^(\d+)번$', line))
        if q_match:
            answer_counter = int(q_match.group(1))
            i += 1
            continue

        if current_section is not None:
            answer_letter = None

            ans_match = re.match(r'정답:\s*([A-D])\)', line)
            if ans_match:
                answer_letter = ans_match.group(1)
                if answer_counter == 0:
                    answer_counter = 1

            if line == '정답' and i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if re.match(r'^[A-D]$', next_line):
                    answer_letter = next_line
                    if answer_counter == 0:
                        answer_counter = 1
                    i += 1

            if answer_letter:
                answer_idx = ord(answer_letter) - ord('A')

                explanation = ''
                j = i + 1
                while j < len(lines):
                    next_line = lines[j]
                    if re.match(r'정답 해설', next_line):
                        exp_text = re.sub(r'^정답 해설:?\s*', '', next_line).strip()
                        if not exp_text and j + 1 < len(lines):
                            j += 1
                            exp_text = lines[j]
                        explanation = exp_text
                        k = j + 1
                        while k < len(lines):
                            if re.match(r'(오답 해설|오답|정답|Question\s+\d+|\d+번|Section\s+\d+|영어 지문)', lines[k]):
                                break
                            explanation += ' ' + lines[k]
                            k += 1
                        break
                    elif re.match(r'(오답|Question\s+\d+|\d+번|Section\s+\d+|영어 지문)', next_line):
                        break
                    j += 1

                key = (current_section, current_module, answer_counter)
                results[key] = {
                    'answer': answer_idx,
                    'answer_letter': answer_letter,
                    'explanation': explanation.strip(),
                }
                answer_counter += 1

        i += 1

    return results


def merge_questions(questions, answers, exam_set):
    merged = []
    # Global question counter across modules
    global_num = 0
    for q in questions:
        global_num += 1
        key = (q['section_num'], q['module_num'], q['question_number'])
        ans = answers.get(key, {})

        section_type = q['section_type'].lower()
        if 'reading' in section_type or 'writing' in section_type:
            section = 'reading-writing'
        else:
            section = section_type.replace(' ', '-')

        merged.append({
            'exam': 'sat',
            'section': section,
            'type': f"s{q['section_num']}m{q['module_num']}",
            'difficulty': 'medium',
            'passage': q.get('passage') or '',
            'question': q['question'],
            'choices': q['choices'],
            'answer': ans.get('answer', -1),
            'explanation': ans.get('explanation', ''),
            'tags': json.dumps(['sat', exam_set, f"s{q['section_num']}m{q['module_num']}"]),
            'examSet': exam_set,
            'questionNumber': global_num,
        })

    return merged


def upload_to_pocketbase(questions):
    token = get_admin_token()
    success = 0
    errors = 0

    for q in questions:
        data = json.dumps({
            'exam': q['exam'],
            'section': q['section'],
            'type': q['type'],
            'difficulty': q['difficulty'],
            'passage': q['passage'],
            'question': q['question'],
            'choices': q['choices'],
            'answer': q['answer'],
            'explanation': q['explanation'],
            'tags': json.loads(q['tags']),
            'examSet': q['examSet'],
            'questionNumber': q['questionNumber'],
        }).encode('utf-8')

        req = urllib.request.Request(
            f"{POCKETBASE_URL}/api/collections/questions/records",
            data=data,
            headers={
                'Content-Type': 'application/json',
                'Authorization': f'Bearer {token}',
            },
            method='POST'
        )

        try:
            with urllib.request.urlopen(req) as resp:
                success += 1
        except urllib.error.HTTPError as e:
            errors += 1
            body = e.read().decode()
            print(f"  Error Q{q['questionNumber']}: {e.code} - {body[:200]}")

    return success, errors


def main():
    if len(sys.argv) < 4:
        print("Usage: python3 scripts/import-sat.py <문제.docx> <해설.docx> <회차이름>")
        print("Example: python3 scripts/import-sat.py q.docx a.docx '2025-sep-cookingsat'")
        sys.exit(1)

    q_path = sys.argv[1]
    a_path = sys.argv[2]
    exam_set = sys.argv[3]

    print(f"Parsing questions: {os.path.basename(q_path)}")
    questions = parse_questions_docx(q_path)
    print(f"  → {len(questions)} questions found")

    print(f"Parsing answers: {os.path.basename(a_path)}")
    answers = parse_explanation_docx(a_path)
    print(f"  → {len(answers)} answers found")

    print(f"Merging with examSet: {exam_set}")
    merged = merge_questions(questions, answers, exam_set)
    answered = sum(1 for m in merged if m['answer'] >= 0)
    print(f"  → {len(merged)} merged ({answered} with answers)")

    if '--dry-run' in sys.argv:
        print("\n[DRY RUN] Sample output:")
        for m in merged[:2]:
            print(json.dumps(m, ensure_ascii=False, indent=2)[:500])
        sys.exit(0)

    print(f"\nUploading to PocketBase ({POCKETBASE_URL})...")
    success, errors = upload_to_pocketbase(merged)
    print(f"  → {success} uploaded, {errors} errors")


if __name__ == "__main__":
    main()
